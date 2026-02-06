from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENTATION
from datetime import datetime
import os

from app.export.word_templates import WORD_district_first_TEMPLATES, WORD_TEMPLATES

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def export_model_to_word(model, processor, court, week):
    template_key = processor.word_template_key
    if not template_key:
        raise ValueError("–£ –ø—Ä–æ—Ü–µ—Å—Å–æ—Ä–∞ –Ω–µ –∑–∞–¥–∞–Ω word_template_key")

    templates = WORD_TEMPLATES.get(template_key)
    if not templates:
        raise ValueError(f"–ù–µ –Ω–∞–π–¥–µ–Ω Word-—à–∞–±–ª–æ–Ω: {template_key}")

    specialization = processor.get_specialization()

    tpl = templates.get(specialization)

    if not tpl:
        raise ValueError(
            f"–ù–µ—Ç Word-—à–∞–±–ª–æ–Ω–∞ –¥–ª—è specialization={specialization}, "
            f"processor={processor.__class__.__name__}"
        )

    document = Document()

    # --- –ê–ª—å–±–æ–º–Ω—ã–π –ª–∏—Å—Ç ---
    section = document.sections[-1]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    # --- –£–∑–∫–∏–µ –ø–æ–ª—è ---
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Inches(0.5))

    # --- –ó–∞–≥–æ–ª–æ–≤–æ–∫ ---
    document.add_paragraph(f"{court} ({specialization}) ‚Äî {week}")

    rows = model.rowCount()
    cols = model.columnCount()

    table = document.add_table(rows=1, cols=cols)
    table.style = "Table Grid"

    # --- –®–∞–ø–∫–∞ ---
    for c in range(cols):
        headers = tpl.get("headers") if tpl else None

        for c in range(cols):
            if headers and c < len(headers):
                table.cell(0, c).text = headers[c]
            else:
                table.cell(0, c).text = model.headerData(c, 1)



    # --- –î–∞–Ω–Ω—ã–µ ---
    for r in range(rows):
        row = table.add_row().cells
        for c in range(cols):
            cell = row[c]
            cell.text = str(model.data(model.index(r, c)))

            p = cell.paragraphs[0]

            # üî• –í–´–†–ê–í–ù–ò–í–ê–ù–ò–ï
            if c == 0:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # –ø–µ—Ä–≤—ã–π —Å—Ç–æ–ª–±–µ—Ü
            else:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # --- –û–±—ä–µ–¥–∏–Ω–µ–Ω–∏—è —Å—Ç–æ–ª–±—Ü–æ–≤ ---
    if tpl:
        for (r1, c1), (r2, c2), text in tpl["merge"]:
            cell = table.cell(r1, c1)
            cell.merge(table.cell(r2, c2))
            cell.text = text

    # –ñ–∏—Ä–Ω–∞—è –ø–µ—Ä–≤–∞—è —Å—Ç—Ä–æ–∫–∞
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
                run.bold = True

    # --- –®—Ä–∏—Ñ—Ç ---
    for row in table.rows:
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)

    # --- –ñ–∏—Ä–Ω–∞—è –ø–æ—Å–ª–µ–¥–Ω—è—è —Å—Ç—Ä–æ–∫–∞ ---
    for cell in table.rows[-1].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # --- –°–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ ---
    filename = f"big_table_{datetime.now():%d.%m.%Y.%H.%M.%S}.docx"
    document.save(filename)
    os.startfile(filename)
