from docx import Document
from docx.shared import Pt
from PyQt5.QtCore import Qt, QModelIndex


def export_table_model_to_word(model, file_path, title=""):
    """
    Экспорт TableModel в Word (.docx)

    model     — QAbstractTableModel (наш TableModel)
    file_path — путь к файлу .docx
    title     — заголовок документа (опционально)
    """

    doc = Document()

    # --- Заголовок ---
    if title:
        heading = doc.add_heading(title, level=1)
        for run in heading.runs:
            run.font.size = Pt(14)

    rows = model.rowCount()
    cols = model.columnCount()

    table = doc.add_table(rows=rows + 1, cols=cols)
    table.style = "Table Grid"

    # --- Шапка таблицы ---
    for col in range(cols):
        cell = table.cell(0, col)
        cell.text = str(model.headerData(col, Qt.Horizontal))

    # --- Данные ---
    for row in range(rows):
        for col in range(cols):
            index = model.index(row, col, QModelIndex())
            value = model.data(index, Qt.DisplayRole)
            table.cell(row + 1, col).text = "" if value is None else str(value)

    doc.save(file_path)
