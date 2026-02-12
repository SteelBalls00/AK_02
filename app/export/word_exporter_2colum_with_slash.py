from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os

from app.export.word_templates import WORD_TEMPLATES


def export_model_to_word(model, processor, court, week):
    # ---------- Проверки ----------
    template_key = processor.word_template_key
    if not template_key:
        raise ValueError("У процессора не задан word_template_key")

    templates = WORD_TEMPLATES.get(template_key)
    if not templates:
        raise ValueError(f"Не найден Word-шаблон: {template_key}")

    specialization = processor.get_specialization()
    tpl = templates.get(specialization)

    if not tpl:
        raise ValueError(
            f"Нет Word-шаблона для specialization={specialization}, "
            f"processor={processor.__class__.__name__}"
        )

    merges = tpl.get("merge", [])

    # ---------- Документ ----------
    document = Document()

    # --- Альбомная ориентация ---
    section = document.sections[-1]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    # --- Узкие поля ---
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Inches(0.5))

    # --- Заголовок ---
    document.add_paragraph(f"{court} ({specialization}) — {week}")

    rows = model.rowCount()
    cols = model.columnCount()

    table = document.add_table(rows=1, cols=cols)
    table.style = "Table Grid"

    # ---------- ШАПКА ----------
    headers = tpl.get("headers")

    for c in range(cols):
        cell = table.cell(0, c)
        if headers and c < len(headers):
            cell.text = headers[c]
        else:
            cell.text = str(model.headerData(c, 1))

        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in p.runs:
            run.font.size = Pt(9)
            run.bold = True

    # ---------- ДАННЫЕ ----------
    for r in range(rows):
        row_cells = table.add_row().cells
        for c in range(cols):
            cell = row_cells[c]
            value = model.data(model.index(r, c))
            cell.text = "" if value is None else str(value)

            p = cell.paragraphs[0]
            if c == 0:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            for run in p.runs:
                run.font.size = Pt(8)

    # ---------- ОБЪЕДИНЕНИЕ ЗАГОЛОВКОВ ----------
    for (r1, c1), (r2, c2), text in merges:
        cell = table.cell(r1, c1)
        cell.merge(table.cell(r2, c2))
        cell.text = text

        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(9)

    # ---------- ОБЪЕДИНЕНИЕ ДАННЫХ + "/" ----------
    header_rows = 1  # у нас одна строка шапки

    for (r1, c1), (r2, c2), _ in merges:
        for r in range(header_rows, header_rows + rows):
            base_cell = table.cell(r, c1)

            values = []
            for c in range(c1, c2 + 1):
                txt = table.cell(r, c).text.strip()
                if txt:
                    values.append(txt)

            merged_text = " / ".join(values)

            base_cell.merge(table.cell(r, c2))
            base_cell.text = merged_text

            p = base_cell.paragraphs[0]
            if c1 == 0:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            for run in p.runs:
                run.font.size = Pt(8)

    # ---------- Жирная последняя строка (ИТОГО) ----------
    for cell in table.rows[-1].cells:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True

    # ---------- Сохранение ----------
    filename = f"big_table_{datetime.now():%d.%m.%Y.%H.%M.%S}.docx"
    document.save(filename)
    os.startfile(filename)
