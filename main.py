# pyinstaller --onedir --noconsole --hidden-import=openpyxl --add-data "Tab_btn.png:." --add-data "Graph_btn.png:." --add-data "Case_analysis.ico:." --icon=Case_analysis.ico --name="Case_analysis" main.py

'''
- путь к базам в файле настроек
- столбцы для бездвижа и возвратов
- закрепить первый столбец с судьями, в случае ширины таблицы за пределы экрана
- в детализации отделить визуально рассмотренные в году
- разделение по категориям????
поправить:

глобальные правки:
- поправить или сделать новый апдейт
'''

import sys
import os
import re
from docx import Document
from datetime import datetime, date
from openpyxl import Workbook
import traceback

from PyQt5.QtWidgets import QFrame, QToolButton, QStackedWidget, QSizePolicy
from PyQt5.QtCore import Qt, QDate, QEasingCurve, QSettings
from PyQt5.QtWidgets import (
    QApplication, QMenu, QMainWindow, QWidget,
    QVBoxLayout, QComboBox, QMessageBox, QTableView,
    QRadioButton, QGroupBox, QHBoxLayout, QPushButton,
    QLabel, QHeaderView, QTextEdit, QSplitter,
    QCalendarWidget, QDialog,
)
from PyQt5.QtGui import QIcon, QFont
from PyQt5.QtCore import QSize, QPropertyAnimation
from PyQt5.QtWidgets import QGraphicsOpacityEffect

from app.constants.pkl_mapping import PKL_MAPPING
from app.repository.bases_repository import BasesRepository
from app.repository.statistics import StatisticsRepository
from app.factory.processor_factory import ProcessorFactory
from app.domain.pkl_selector import select_pkl_for_context
from app.ui.table_model import TableModel
from app.workers.data_load_worker import DataLoadWorker
from app.export.word_exporter import export_model_to_word
from app.ui.graph_widget import GraphWidget


BASE_DIR = os.path.join(os.path.dirname(__file__), "bases")


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Анализ картотек")

        script_dir = os.path.dirname(os.path.abspath(__file__))
        icon_path = os.path.join(script_dir, "Case_analysis.ico")
        self.setWindowIcon(QIcon(icon_path))

        self.bases_repo = BasesRepository(BASE_DIR)
        self.stats_repo = StatisticsRepository()

        # ====== СОСТОЯНИЕ (ДО UI!) ======
        self.specialization = "GPK"
        self.instance = "first"

        self.current_pkl_path = None
        self.current_raw_data = None
        self.current_context = None

        self.week_index = 0
        self.max_week_index = 0
        self.current_week_key = None

        self.active_workers = []

        # self.settings = QSettings("CaseAnalysis", "CaseAnalysisApp")
        self.settings = QSettings("settings.ini", QSettings.IniFormat)
        self.settings.setIniCodec("UTF-8")

        # ====== UI ======
        self._init_ui()
        self._load_courts()

    def _init_ui(self):
        self._ui_ready = False

        central = QWidget(self)
        main_layout = QVBoxLayout(central)

        # ================= Верхняя панель =================
        header_widget = QWidget()
        header_widget.setObjectName("panel")
        top_layout = QHBoxLayout(header_widget)
        top_layout.setContentsMargins(8, 4, 8, 4)
        top_layout.setSpacing(6)

        # --- Переключение недель ---
        self.prev_week_btn = QPushButton("◀")
        self.next_week_btn = QPushButton("▶")
        self.week_label = QLabel("")
        self.week_label.setAlignment(Qt.AlignCenter)
        self.week_label.setCursor(Qt.PointingHandCursor)
        self.week_label.mousePressEvent = self.on_week_label_clicked

        self.week_label.setProperty("role", "week-label")
        self.prev_week_btn.setProperty("role", "week-nav")
        self.next_week_btn.setProperty("role", "week-nav")

        self.prev_week_btn.clicked.connect(self.prev_week)
        self.next_week_btn.clicked.connect(self.next_week)

        for btn in (self.prev_week_btn, self.next_week_btn):
            btn.setFixedSize(68, 48)

        self.header_stack = QStackedWidget()
        self.header_stack.setSizePolicy(
            QSizePolicy.Maximum,
            QSizePolicy.Fixed
        )
        self.header_stack.setFixedHeight(80)
        top_layout.addWidget(self.header_stack)

        self.week_nav_widget = QWidget()
        week_layout = QHBoxLayout(self.week_nav_widget)
        week_layout.setContentsMargins(0, 0, 0, 0)

        week_layout.addWidget(self.prev_week_btn)
        week_layout.addWidget(self.week_label)
        week_layout.addWidget(self.next_week_btn)
        week_layout.addStretch()

        self.header_stack.addWidget(self.week_nav_widget)

        # --- Суд ---
        self.court_group = QGroupBox("Суд")
        court_layout = QVBoxLayout(self.court_group)

        self.court_combo = QComboBox()
        self.court_combo.currentTextChanged.connect(self.on_court_changed)

        court_layout.addWidget(self.court_combo)

        top_layout.addWidget(self.court_group)

        # --- Специализация ---
        spec_group = QGroupBox("Специализация")
        self.spec_layout = QHBoxLayout(spec_group)

        self.spec_buttons = {}

        specs = {
            "GPK": "ГПК",
            "KAS": "КАС",
            "AP": "АП",
            "AP1": "АП1",
            "U1": "УГ",
            "M_U1": "М.Уг",
            "M_AOS": "М.",
        }

        for code, label in specs.items():
            rb = QRadioButton(label)
            rb.setProperty("spec", code)
            rb.toggled.connect(self.on_context_changed)
            self.spec_buttons[code] = rb
            self.spec_layout.addWidget(rb)

        self.spec_buttons["GPK"].setChecked(True)
        top_layout.addWidget(spec_group)

        # --- Инстанция ---
        inst_group = QGroupBox("Инстанция")
        inst_layout = QVBoxLayout(inst_group)

        self.instance_buttons = {}

        first_btn = QRadioButton("1 инстанция")
        first_btn.instance_value = "first"

        appeal_btn = QRadioButton("Апелляция")
        appeal_btn.instance_value = "appeal"

        self.instance_buttons["first"] = first_btn
        self.instance_buttons["appeal"] = appeal_btn

        appeal_btn.toggled.connect(self.on_context_changed)
        first_btn.toggled.connect(self.on_context_changed)

        inst_layout.addWidget(first_btn)
        inst_layout.addWidget(appeal_btn)

        self.instance_buttons["first"].setChecked(True)
        top_layout.addWidget(inst_group)

        # --- Кнопка выгрузки в Word ---
        script_dir = os.path.dirname(os.path.abspath(__file__))
        word_icon_path = os.path.join(script_dir, "Word_png.png")

        script_dir = os.path.dirname(os.path.abspath(__file__))

        self.view_table_btn = QToolButton()
        self.view_chart_btn = QToolButton()

        self.view_table_btn.setIcon(QIcon(os.path.join(script_dir, "Tab_btn.png")))
        self.view_chart_btn.setIcon(QIcon(os.path.join(script_dir, "Graph_btn.png")))

        # размер иконки (уменьшили)
        self.view_table_btn.setIconSize(QSize(90, 60))
        self.view_chart_btn.setIconSize(QSize(90, 60))

        # текст
        self.view_table_btn.setText("Таблица")
        self.view_chart_btn.setText("График")

        # иконка сверху, текст снизу
        self.view_table_btn.setToolButtonStyle(Qt.ToolButtonTextUnderIcon)
        self.view_chart_btn.setToolButtonStyle(Qt.ToolButtonTextUnderIcon)

        for btn in (self.view_table_btn, self.view_chart_btn):
            btn.setCheckable(True)
            btn.setAutoExclusive(True)  # чтобы работали как вкладки
            btn.setCursor(Qt.PointingHandCursor)
            btn.setObjectName("viewSwitchBtn")
            btn.setFixedSize(110, 95)

        self.view_table_btn.setChecked(True)

        self.view_table_btn.clicked.connect(self.switch_to_table)
        self.view_chart_btn.clicked.connect(self.switch_to_chart)

        top_layout.addWidget(self.view_table_btn)
        top_layout.addWidget(self.view_chart_btn)

        self.word_export_btn = QPushButton()
        self.word_export_btn.setIcon(QIcon(word_icon_path))
        self.word_export_btn.setIconSize(QSize(86, 25))
        self.word_export_btn.clicked.connect(self.export_to_word)
        self.word_export_btn.setObjectName("export_to_word")

        top_layout.addWidget(self.word_export_btn)

        # растяжка, чтобы элементы не слипались
        top_layout.addStretch()

        # ================= переключатель темы =================
        self.theme_toggle_btn = QToolButton()
        self.theme_toggle_btn.setText("🌙 Тёмная тема")
        self.theme_toggle_btn.setCheckable(True)
        self.theme_toggle_btn.setChecked(False)  # стартуем с тёмной

        self.theme_toggle_btn.clicked.connect(self.toggle_theme)

        top_layout.addWidget(self.theme_toggle_btn)

        # ================= Черточка перед таблицей =================
        self.splitter = QSplitter(Qt.Vertical)

        # ================= Таблица =================
        self.table_view = QTableView()
        self.model = TableModel()
        self.table_view.setModel(self.model)
        self.table_view.setSortingEnabled(True)

        self.table_view.setAlternatingRowColors(True)
        self.table_view.setShowGrid(True)
        self.table_view.verticalHeader().setVisible(False)
        self.table_view.horizontalHeader().setStretchLastSection(True)

        '''
        Настройка высоты строк
        20–22 — очень компактно
        24 — комфортно
        28 — «воздушно»
        '''

        vertical_header = self.table_view.verticalHeader()
        vertical_header.setDefaultSectionSize(24)  # Настройка высоты строк

        self.table_opacity = QGraphicsOpacityEffect(self.table_view.viewport())
        self.table_view.viewport().setGraphicsEffect(self.table_opacity)
        self.table_opacity.setOpacity(1.0)  # ВАЖНО

        self.fade_anim = QPropertyAnimation(self.table_opacity, b"opacity")
        self.fade_anim.setDuration(150)
        self.fade_anim.setEasingCurve(QEasingCurve.InOutQuad)

        header = self.table_view.horizontalHeader()
        header.setSectionResizeMode(QHeaderView.ResizeToContents)
        header.setDefaultAlignment(Qt.AlignCenter)

        # не обрезать текст троеточием
        header.setTextElideMode(Qt.ElideNone)  # Управляет обрезкой текста, если он не влезает

        # центрирование
        header.setDefaultAlignment(Qt.AlignCenter)  # Центрирует текст внутри ячеек заголовка

        # ширина под содержимое
        header.setSectionResizeMode(QHeaderView.ResizeToContents)  # ширина столбца = ширина самого широкого содержимого

        # даём место для многострочных заголовков
        header.setFixedHeight(70)

        self.table_view.setStyleSheet("""
        QHeaderView::section {
            padding: 6px;
            font-weight: bold;
        }
        """)

        # --- Область детализации ---
        self.details_view = QTextEdit()

        self.details_view.setReadOnly(True)
        self.details_view.setMinimumHeight(30)
        self.details_view.setContextMenuPolicy(Qt.CustomContextMenu)
        self.details_view.customContextMenuRequested.connect(
            self.show_details_context_menu
        )
        self.details_view.setLineWrapMode(QTextEdit.WidgetWidth)
        self.details_view.setFont(QFont("Consolas", 10))

        self.details_view.setPlaceholderText(
            "Выберите ячейку таблицы, чтобы увидеть детализацию"
        )

        selection_model = self.table_view.selectionModel()
        selection_model.selectionChanged.connect(self.on_table_selection_changed)

        # ================= Разделитель =================
        separator = QFrame()
        separator.setFrameShape(QFrame.HLine)
        separator.setFrameShadow(QFrame.Sunken)

        # ================= Сборка =================
        self.stacked_widget = QStackedWidget()
        self.stacked_widget.addWidget(self.table_view)

        self.graph_widget = GraphWidget()
        self.graph_widget.point_clicked.connect(self.on_graph_point_clicked)

        # ===== begin Управление датами графика =====
        self.date_group = QGroupBox("Диапазон дат для графика")
        self.date_group.setSizePolicy(
            QSizePolicy.Maximum,
            QSizePolicy.Fixed
        )
        date_layout = QHBoxLayout(self.date_group)
        date_layout.setContentsMargins(8, 4, 8, 4)
        date_layout.setSpacing(6)


        self.chart_date_from = self.graph_widget.date_from
        self.chart_date_to = self.graph_widget.date_to
        self.chart_date_from.setFixedWidth(200)
        self.chart_date_to.setFixedWidth(200)

        # увеличиваем размер
        # self.chart_date_from.setMinimumHeight(32)
        # self.chart_date_to.setMinimumHeight(32)

        font = self.chart_date_from.font()
        font.setPointSize(font.pointSize() + 1)
        self.chart_date_from.setFont(font)
        self.chart_date_to.setFont(font)

        date_layout.addWidget(QLabel("С:"))
        date_layout.addWidget(self.chart_date_from)
        date_layout.addSpacing(10)
        date_layout.addWidget(QLabel("По:"))
        date_layout.addWidget(self.chart_date_to)
        date_layout.addStretch()

        self.header_stack.addWidget(self.date_group)
        # ===== end Управление датами графика =====

        self.stacked_widget.addWidget(self.graph_widget)

        # график добавим позже
        self.splitter.addWidget(self.stacked_widget)
        self.splitter.addWidget(self.details_view)
        self.splitter.setStretchFactor(0, 8)  # таблица
        self.splitter.setStretchFactor(1, 6)  # детализация
        self.splitter.setSizes([700, 300])

        main_layout.addWidget(header_widget)
        main_layout.addWidget(separator)
        main_layout.addWidget(self.splitter)

        self.setCentralWidget(central)

        self._ui_ready = True

    def closeEvent(self, event):
        self.settings.setValue("court", self.court_combo.currentText())
        self.settings.setValue("specialization", self.specialization)
        self.settings.setValue("instance", self.instance)
        event.accept()

    def switch_to_table(self):
        self.view_table_btn.setChecked(True)
        self.view_chart_btn.setChecked(False)
        self.stacked_widget.setCurrentIndex(0)
        self.header_stack.setCurrentIndex(0)

    def switch_to_chart(self):
        self.view_chart_btn.setChecked(True)
        self.view_table_btn.setChecked(False)
        self.stacked_widget.setCurrentIndex(1)
        self.header_stack.setCurrentIndex(1)

    def set_radio_visible(self, btn, visible: bool):
        if not visible and btn.isChecked():
            btn.setAutoExclusive(False)
            btn.setChecked(False)
            btn.setAutoExclusive(True)

        btn.setVisible(visible)

    def update_instance_buttons(self, court_name: str):
        instances = self.bases_repo.get_available_instances(
            court_name,
            self.specialization
        )

        for inst, btn in self.instance_buttons.items():
            btn.setEnabled(inst in instances)

        # защита от невалидного состояния
        if self.instance not in instances and instances:
            self.instance = next(iter(instances))
            self.instance_buttons[self.instance].setChecked(True)

    def update_specialization_buttons(self, court_name: str):
        available_specs = self.bases_repo.get_available_specializations(court_name, self.instance)

        for spec, btn in self.spec_buttons.items():
            self.set_radio_visible(btn, spec in available_specs)

        # гарантируем выбранную специализацию
        if self.specialization not in available_specs and available_specs:
            new_spec = next(iter(available_specs))
            self.spec_buttons[new_spec].setChecked(True)
            self.specialization = new_spec

    def animate_table_update(self, update_callback):
        """
        Полностью безопасное обновление таблицы:
        - без мигания
        - без микро-дёрганий
        - без призраков старых данных
        """

        # если анимация уже идёт — остановить
        if self.fade_anim.state() == QPropertyAnimation.Running:
            self.fade_anim.stop()

        # 1. МГНОВЕННО скрываем содержимое таблицы
        self.table_opacity.setOpacity(0.0)

        # 2. Полностью блокируем перерисовку
        self.table_view.setUpdatesEnabled(False)
        # 3. Применяем данные
        update_callback()
        # 4. Разрешаем перерисовку
        self.table_view.setUpdatesEnabled(True)

        # 5. Плавно показываем новую таблицу
        self.fade_anim.setStartValue(0.0)
        self.fade_anim.setEndValue(1.0)
        self.fade_anim.start()

    def toggle_theme(self, checked: bool):
        '''Переключение цвета темы'''
        app = QApplication.instance()

        if checked:
            app.setStyleSheet(DARK_STYLE)
            self.theme_toggle_btn.setText("🌞 Светлая тема")
            self.graph_widget.apply_dark_style()
            self.graph_widget.update_chart()
        else:
            app.setStyleSheet(LIGHT_STYLE)
            self.theme_toggle_btn.setText("🌙 Тёмная тема")
            self.graph_widget.apply_light_style()
            self.graph_widget.update_chart()

    def select_week_by_date(self, selected_date: date):
        """
        Выбирает неделю, в которую попадает дата.
        Если такой нет — выбирает ближайшую.
        """
        weeks = list(self.current_raw_data.keys())

        parsed_weeks = []

        for idx, week_str in enumerate(weeks):
            try:
                start_str, end_str = week_str.split(" - ")
                start = datetime.strptime(start_str, "%d.%m.%Y").date()
                end = datetime.strptime(end_str, "%d.%m.%Y").date()
                parsed_weeks.append((idx, start, end))
            except Exception:
                continue

        if not parsed_weeks:
            return

        # 1️⃣ Пытаемся найти точное попадание
        for idx, start, end in parsed_weeks:
            if start <= selected_date <= end:
                self.week_index = idx
                self.reload_current_court()
                return

        # 2️⃣ Ищем ближайшую неделю
        def distance(week):
            _, start, end = week
            if selected_date < start:
                return (start - selected_date).days
            if selected_date > end:
                return (selected_date - end).days
            return 0

        closest = min(parsed_weeks, key=distance)
        self.week_index = closest[0]
        self.reload_current_court()

    def on_calendar_confirmed(self, calendar: QCalendarWidget, dialog: QDialog):
        qdate = calendar.selectedDate()
        selected_date = date(qdate.year(), qdate.month(), qdate.day())

        self.select_week_by_date(selected_date)

        dialog.accept()

    def on_week_label_clicked(self, event):
        if not self.current_raw_data:
            return

        dialog = QDialog(self)
        dialog.setWindowTitle("Выбор даты")
        dialog.setModal(True)

        layout = QVBoxLayout(dialog)

        calendar = QCalendarWidget()
        calendar.setGridVisible(True)
        calendar.setSelectedDate(QDate.currentDate())

        layout.addWidget(calendar)

        btn_ok = QPushButton("Выбрать")
        layout.addWidget(btn_ok)

        btn_ok.clicked.connect(lambda: self.on_calendar_confirmed(calendar, dialog))

        dialog.resize(300, 250)
        dialog.exec_()

    def parse_details_blocks(self):
        """
        Разбивает детализацию на блоки для экспорта,
        сохраняя ВСЕ строки внутри блока в исходном порядке.

        Новый блок начинается с:
        - "Судья:"
        - "Неделя:"
        """

        text = self.details_view.toPlainText()
        lines = [l.rstrip() for l in text.splitlines()]

        blocks = []
        current_block = []

        def is_block_start(line: str) -> bool:
            return line.startswith("Судья:") or line.startswith("Неделя:")

        for line in lines:
            if not line.strip():
                continue

            if is_block_start(line) and current_block:
                blocks.append(current_block)
                current_block = []

            current_block.append(line)

        if current_block:
            blocks.append(current_block)

        return blocks

    def export_details_to_excel(self, only_numbers: bool):
        blocks = self.parse_details_blocks()

        wb = Workbook()
        ws = wb.active
        ws.title = "Детализация"

        row = 1

        for block in blocks:
            for line in block:
                value = line

                if only_numbers and line.strip().startswith("•"):
                    value = self.extract_case_number(line)

                ws.cell(row=row, column=1, value=value)
                row += 1

            row += 2  # пустая строка между блоками

        filename = f"details_{datetime.now():%d.%m.%Y.%H.%M.%S}.xlsx"
        wb.save(filename)
        os.startfile(filename)

    def export_details_to_word(self, only_numbers: bool):
        blocks = self.parse_details_blocks()

        def get_judge_name(block):
            for line in block:
                if line.startswith("Судья:"):
                    return line.replace("Судья:", "", 1).strip()
            return "Без судьи"

        # группируем блоки по судье
        grouped_by_judge = {}
        judge_order = []

        for block in blocks:
            judge = get_judge_name(block)
            if judge == 'Без судьи':
                continue
            if judge not in grouped_by_judge:
                grouped_by_judge[judge] = []
                judge_order.append(judge)
            grouped_by_judge[judge].append(block)

        document = Document()
        # document.add_heading("Детализация", level=1)

        for judge_index, judge in enumerate(judge_order):
            judge_blocks = grouped_by_judge[judge]

            # заголовок страницы
            document.add_paragraph(f"Судья: {judge}")
            document.add_paragraph("")

            for block in judge_blocks:
                for line in block:
                    # строку "Судья: ..." внутри блока повторно не печатаем
                    if line.startswith("Судья:"):
                        continue

                    value = line
                    if only_numbers and line.strip().startswith("•"):
                        value = self.extract_case_number(line)

                    document.add_paragraph(value)

                document.add_paragraph("")  # отступ между показателями одного судьи

            # новая страница только между судьями
            if judge_index < len(judge_order) - 1:
                document.add_page_break()

        filename = f"details_{datetime.now():%d.%m.%Y.%H.%M.%S}.docx"
        document.save(filename)
        os.startfile(filename)

    def get_details_lines(self):
        """
        Возвращает список строк детализации (без пустых)
        """
        text = self.details_view.toPlainText()
        return [line.strip() for line in text.splitlines() if line.strip()]

    def extract_case_number(self, line: str) -> str:
        """
        Извлекает номер дела до первой запятой
        """
        if "," in line:
            return line.split(",", 1)[0].strip().replace('• ', '')
        return line.strip()

    def copy_details_to_clipboard(self):
        blocks = self.parse_details_blocks()

        lines = []
        for block in blocks:
            lines.extend(block)
            lines.append("")

        QApplication.clipboard().setText("\n".join(lines))

    def show_details_context_menu(self, pos):
        menu = QMenu(self)

        # --- Копировать ---
        copy_action = menu.addAction("Скопировать в буфер обмена")
        copy_action.triggered.connect(self.copy_details_to_clipboard)

        menu.addSeparator()

        # --- Word ---
        word_menu = menu.addMenu("Передать в Word")
        word_only_numbers = word_menu.addAction("Только номера дел")
        word_full = word_menu.addAction("Номера дел со всей информацией")

        word_only_numbers.triggered.connect(
            lambda: self.export_details_to_word(only_numbers=True)
        )
        word_full.triggered.connect(
            lambda: self.export_details_to_word(only_numbers=False)
        )

        # --- Excel ---
        excel_menu = menu.addMenu("Передать в Excel")
        excel_only_numbers = excel_menu.addAction("Только номера дел")
        excel_full = excel_menu.addAction("Номера дел со всей информацией")

        excel_only_numbers.triggered.connect(
            lambda: self.export_details_to_excel(only_numbers=True)
        )
        excel_full.triggered.connect(
            lambda: self.export_details_to_excel(only_numbers=False)
        )

        menu.exec_(self.details_view.mapToGlobal(pos))



    def _format_details_block(self, judge, column, details):


        def normalize_case_line(raw: str) -> str:
            _PREFIX_RE = re.compile(r"\d\.\d{3}-")
            """
            Удаляет ТОЛЬКО префикс вида '2.123-' (цифра + точка + 3 цифры + дефис).
            Если такого шаблона нет — строка возвращается без изменений.
            """
            return _PREFIX_RE.sub("", raw, count=1)

        column = column.replace('\n', ' ')
        if column != 'Судья':
            lines = [
                f"Судья: {judge}",
                f"Показатель: {column}",
            ]

        if not details:
            lines.append("Детализация отсутствует")
            return "\n".join(lines)


        for title, values in details:
            total = len(values) if values else 0
            lines.append(f"{title}: {total}")

            for v in values:
                v = normalize_case_line(v)
                lines.append(f"  • {v}")
            lines.append(f'{"-"*60}')

        return "\n".join(lines)

    def on_table_selection_changed(self, selected, deselected):
        def normalize_case_line(raw: str) -> str:
            _PREFIX_RE = re.compile(r"\d\.\d{3}-")
            """
            Удаляет ТОЛЬКО префикс вида '2.123-' (цифра + точка + 3 цифры + дефис).
            Если такого шаблона нет — строка возвращается без изменений.
            """
            return _PREFIX_RE.sub("", raw, count=1)

        if not self.current_context:
            return

        indexes = self.table_view.selectionModel().selectedIndexes()
        if not indexes:
            self.details_view.clear()
            return

        blocks = []

        for index in indexes:
            row = index.row()
            col = index.column()

            judge_name = self.model.data(self.model.index(row, 0))
            column_name = self.model.headerData(col, Qt.Horizontal)

            # ---- ЕСЛИ строка "Всего"
            if judge_name == "Всего" and col != 0:

                week_key = list(self.current_raw_data.keys())[self.week_index]
                week_data = self.current_raw_data.get(week_key, {})

                lines = [
                    f"Неделя: {week_key}",
                    f"Показатель: {column_name}",
                    "Всего по всем судьям",
                    ""
                ]

                judges_data = []

                # --- собираем данные по всем судьям
                for judge in week_data.keys():

                    details = self.current_processor.get_cell_details(
                        judge=judge,
                        column=column_name,
                        week_index=self.week_index,
                    )

                    def extract_case_key(raw: str) -> str:
                        raw = raw.strip()
                        if "," in raw:
                            return raw.split(",", 1)[0].strip()
                        return raw

                    unique_cases = set()

                    for _, values in details:
                        for v in values:
                            unique_cases.add(extract_case_key(v))

                    total_cases = len(unique_cases)

                    if total_cases > 0:
                        judges_data.append((judge, total_cases, details))

                # 🔥 сортировка по убыванию количества дел
                judges_data.sort(key=lambda x: x[1], reverse=True)

                if not judges_data:
                    self.details_view.setPlainText("Детализация отсутствует.")
                    return

                # --- вывод
                for judge, total_cases, details in judges_data:

                    lines.append(f"Судья: {judge} — дел: {total_cases}")

                    for title, values in details:
                        lines.append(f"{title}: {len(values)}")

                        for v in values:
                            lines.append(f"  • {normalize_case_line(v)}")

                    lines.append("-" * 40)
                    lines.append("")

                self.details_view.setPlainText("\n".join(lines))

            else:
                # обычная логика для судьи
                details = self.current_processor.get_cell_details(
                    judge=judge_name,
                    column=column_name,
                    week_index=self.week_index,
                )

                blocks.append(self._format_details_block(
                    judge_name, column_name, details
                ))

        if blocks:
            self.details_view.setPlainText("\n\n".join(blocks))

    def restore_last_selection(self, courts):
        saved_court = self.settings.value("court")
        saved_spec = self.settings.value("specialization")
        saved_instance = self.settings.value("instance")

        restored = False

        if saved_court and saved_court in courts:
            index = self.court_combo.findText(saved_court)
            if index >= 0:
                self.court_combo.setCurrentIndex(index)
                restored = True

        if saved_spec and saved_spec in self.spec_buttons:
            self.spec_buttons[saved_spec].setChecked(True)
            self.specialization = saved_spec

        if saved_instance and saved_instance in self.instance_buttons:
            self.instance_buttons[saved_instance].setChecked(True)
            self.instance = saved_instance

        return restored

    def _load_courts(self):
        courts = self.bases_repo.get_courts_with_any_pkls()

        self.court_combo.clear()
        self.court_combo.addItems(courts)

        if courts:
            restored = self.restore_last_selection(courts)

            if not restored:
                self.court_combo.setCurrentIndex(0)

            # 🔥 ВАЖНО — принудительно загружаем
            self.on_court_changed(self.court_combo.currentText())

        # показать / скрыть groupbox
        self.court_group.setVisible(len(courts) > 1)

    def reload_current_court(self):
        if not hasattr(self, "court_combo"):
            return

        court = self.court_combo.currentText()
        if court:
            self.on_court_changed(court)

    def on_context_changed(self):
        if not getattr(self, "_ui_ready", False):
            return

        # specialization
        for spec, btn in self.spec_buttons.items():
            if btn.isChecked():
                self.specialization = spec
                break

        # instance
        for inst, btn in self.instance_buttons.items():
            if btn.isChecked():
                self.instance = inst
                break

        self.reload_current_court()

    def on_court_changed(self, court_name):
        if not court_name:
            return

        pkl_files = self.bases_repo.get_pkl_files(court_name)

        if not pkl_files:
            self.model.set_table_data({})
            return

        # =========================================
        # 1️⃣ Проверяем — существует ли текущая комбинация
        # =========================================
        pkl_name = select_pkl_for_context(
            pkl_files,
            specialization=self.specialization,
            instance=self.instance
        )

        # =========================================
        # 2️⃣ Если нет — автоматически ищем валидную
        # =========================================
        if not pkl_name:
            found = False

            for spec in self.spec_buttons.keys():
                for inst in ("first", "appeal"):
                    candidate = select_pkl_for_context(
                        pkl_files,
                        specialization=spec,
                        instance=inst
                    )
                    if candidate:
                        self.specialization = spec
                        self.instance = inst

                        self.spec_buttons[spec].setChecked(True)
                        self.instance_buttons[inst].setChecked(True)

                        pkl_name = candidate
                        found = True
                        break
                if found:
                    break

        # если вообще нет ни одной подходящей базы
        if not pkl_name:
            QMessageBox.warning(
                self,
                "Нет данных",
                "Для выбранного суда нет подходящей базы"
            )
            self.model.set_table_data({})
            return

        # =========================================
        # 3️⃣ Обновляем доступность кнопок UI
        # =========================================
        self.update_specialization_buttons(court_name)
        self.update_instance_buttons(court_name)

        # =========================================
        # 4️⃣ Загружаем pkl
        # =========================================
        pkl_path = self.bases_repo.get_pkl_path(court_name, pkl_name)

        # если тот же файл — просто обновляем таблицу
        if self.current_pkl_path == pkl_path and self.current_raw_data is not None:
            self.load_table_async()
            return

        raw_data, context = self.stats_repo.load(pkl_path)

        self.current_raw_data = raw_data
        self.current_context = context
        self.current_pkl_path = pkl_path

        # обновляем график
        self.graph_widget.set_data(
            raw_data=self.current_raw_data,
            processor=ProcessorFactory.get(context)
        )

        weeks = list(raw_data.keys())
        self.max_week_index = max(0, len(weeks) - 1)

        # пытаемся сохранить текущую неделю
        if self.current_week_key in weeks:
            self.week_index = weeks.index(self.current_week_key)
        else:
            self.week_index = self.max_week_index

        self.load_table_async()

        if self.week_index > self.max_week_index:
            self.week_index = self.max_week_index

        self.table_view.resizeColumnsToContents()

    def load_table_async(self):
        self.table_view.setEnabled(False)

        # 1. Получаем процессор из фабрики
        processor = ProcessorFactory.get(self.current_context)

        # 2. Сохраняем его для детализации
        self.current_processor = processor

        # 3. Запускаем воркер
        worker = DataLoadWorker(
            processor=processor,
            raw_data=self.current_raw_data,
            week_index=self.week_index
        )

        self.active_workers.append(worker)

        worker.finished.connect(
            lambda table_data, w=worker: self.on_data_loaded(table_data, w)
        )
        worker.error.connect(self.on_data_error)

        worker.start()

    def on_graph_point_clicked(self, data):
        def normalize_case_line(raw: str) -> str:
            _PREFIX_RE = re.compile(r"\d\.\d{3}-")
            """
            Удаляет ТОЛЬКО префикс вида '2.123-' (цифра + точка + 3 цифры + дефис).
            Если такого шаблона нет — строка возвращается без изменений.
            """
            return _PREFIX_RE.sub("", raw, count=1)

        week_key = data["week_key"]
        category = data["category"]
        is_double = data["double_click"]

        weeks = list(self.current_raw_data.keys())

        if week_key not in weeks:
            return

        real_week_index = weeks.index(week_key)

        # двойной клик → перейти к таблице
        if is_double:
            self.week_index = real_week_index
            self.switch_to_table()
            self.reload_current_court()
            return

        week_data = self.current_raw_data.get(week_key, {})

        lines = [
            f"Неделя: {week_key}",
            f"Показатель: {category}",
            ""
        ]

        # ===================================================
        # 🔥 РЕЖИМ СРАВНЕНИЯ КАТЕГОРИЙ
        # ===================================================
        if self.graph_widget.compare_mode.isChecked():

            judges_with_counts = []

            for judge, judge_data in week_data.items():
                cases = judge_data.get(category, [])
                count = len(cases)

                if count > 0:
                    judges_with_counts.append((judge, count))

            # сортировка по убыванию
            judges_with_counts.sort(key=lambda x: x[1], reverse=True)

            if not judges_with_counts:
                self.details_view.setPlainText("Детализация отсутствует.")
                return

            for judge, count in judges_with_counts:
                lines.append(f"Судья: {judge} — дел: {count}")

            self.details_view.setPlainText("\n".join(lines))
            return

        # ===================================================
        # 🔥 ОБЫЧНЫЙ РЕЖИМ (СУДЬИ)
        # ===================================================

        judges = data["judges"]

        has_data = False

        for judge in judges:
            judge_data = week_data.get(judge, {})
            cases = judge_data.get(category, [])

            if not cases:
                continue

            has_data = True

            lines.append(f"Судья: {judge} — дел: {len(cases)}")

            for case in cases:
                lines.append(f"  • {normalize_case_line(case)}")

            lines.append("-" * 40)
            lines.append("")

        if not has_data:
            self.details_view.setPlainText("Детализация отсутствует.")
        else:
            self.details_view.setPlainText("\n".join(lines))

    def on_data_loaded(self, table_data, worker):
        def apply():
            self.model.set_table_data(table_data)

            # сортировка по судье
            self.table_view.sortByColumn(0, Qt.AscendingOrder)

            # неделя
            self.week_label.setText(table_data.get("week", ""))

            # UI
            self.table_view.setEnabled(True)

            self.current_week_key = table_data.get("week")

        self.animate_table_update(apply)

        if worker in self.active_workers:
            self.active_workers.remove(worker)

    def on_data_error(self, message, worker):
        QMessageBox.critical(self, "Ошибка загрузки", message)
        self.table_view.setEnabled(True)

        if worker in self.active_workers:
            self.active_workers.remove(worker)

    def prev_week(self):
        if self.week_index > 0:
            self.week_index -= 1
            self.reload_current_court()

    def next_week(self):
        if self.week_index < self.max_week_index:
            self.week_index += 1
            self.reload_current_court()

    def export_to_word(self):
        if self.model.rowCount() == 0:
            QMessageBox.information(self, "Нет данных", "Таблица пуста")
            return

        export_model_to_word(
            model=self.model,
            processor=self.current_processor,
            court=self.court_combo.currentText(),
            week=self.week_label.text()
        )

LIGHT_STYLE = """
QWidget {
    font-family: "Segoe UI";
    color: #2b2b2b;
}

/* --- Таблица --- */
QTableView {
    background-color: #ffffff;
    gridline-color: #dcdcdc;
    selection-background-color: #e6f0fa;
    selection-color: #000000;
    alternate-background-color: #fafafa;
}

QTableView::item:selected {
    background-color: #cfe3f6;
}
/* ================ QPushButton ================ */
QPushButton {
    background-color: #3a6ea5;
    color: white;
    border: none;
    padding: 6px 12px;
    border-radius: 4px;
}

QPushButton:hover {
    background-color: #4a82c0;
}

QPushButton:pressed {
    background-color: #2f5d8a;
}

QPushButton#export_to_word {
    background-color: transparent;
    min-width: 48px;
    min-height: 48px;
    padding: 0px;
}

QPushButton#export_to_word:hover {
    background-color: #5a96d5;
}

QPushButton[role="week-nav"] {
    background-color: #4a86c5;
    border: none;
    font-weight: bold;
    padding: 4px 8px;
    font-size: 20pt;
    font-weight: bold;
}

QPushButton[role="week-nav"]:hover {
    background-color: #4a86c5;
    font-size: 20pt;
    font-weight: bold;
}

QPushButton[role="week-nav"]:pressed {
    background-color: rgba(0, 0, 0, 0.15);
}
QToolButton#viewSwitchBtn {
    background: transparent;
    border: none;
    padding: 4px;
    font-weight: bold;
}

QToolButton#viewSwitchBtn:hover {
    background-color: rgba(74,134,197,0.08);
    border-radius: 12px;
}

QToolButton#viewSwitchBtn:checked {
    background-color: rgba(74,134,197,0.18);
    border-radius: 14px;
}

/* ================ RadioButton ================ */
/* --- Radio / Check --- */
QRadioButton, QCheckBox {
    spacing: 6px;
    font-weight: bold;
}
QRadioButton[spec="GPK"] { font-weight: bold; }

QRadioButton::indicator {
    width: 14px;
    height: 14px;
}

QRadioButton::indicator:disabled {
    background-color: #c0c0c0;
    width: 14px;
    height: 14px;
    border-radius: 7px;
}

/* ================== LABEL ================== */
QLabel {
    color: black;
}

QLabel[role="week-label"] {
    font-size: 20pt;
    font-weight: bold;
}
/* ================== ComboBox ================== */
QComboBox {
    background-color: #ffffff;
    border: 1px solid #cfcfcf;
    padding: 4px;
    border-radius: 4px;
    min-height: 18px;
}
/* ================== QTextEdit ================== */
QTextEdit {
    background-color: #fcfcfc;
    border: 1px solid #cfcfcf;
    border-radius: 4px;
    padding: 6px;
}
/* ================== ToolButton ================== */
QToolButton {
    background-color: transparent;
    border: none;
    padding: 4px;
}

QToolButton:hover {
    background-color: #e6f0fa;
}
/* ================== QGroupBox ================== */
QGroupBox {
    font-weight: bold;
    border: 1px solid #cfcfcf;
    border-radius: 6px;
    margin-top: 8px;
}

QGroupBox::title {
    subcontrol-origin: margin;
    left: 10px;
    padding: 0 4px;
}
/* ================== QDateEdit ================== */
QDateEdit {
    background-color: #ffffff;
    border: 1px solid #cfcfcf;
    border-radius: 4px;
    padding: 4px 6px;
    min-height: 18px;
}

QDateEdit:hover {
    border: 1px solid #a8c5e6;
}

QDateEdit:focus {
    border: 1px solid #4a86c5;
}


QDateEdit::down-arrow {
    width: 10px;
    height: 10px;
}
"""

DARK_STYLE = """
/* ================== БАЗА ================== */
QWidget {
    background-color: #2b2b2b;
    color: #e6e6e6;
    font-family: "Segoe UI";
}

/* ================== ПАНЕЛИ ================== */
QFrame#panel {
    background-color: #313335;
    border: 1px solid #444444;
    border-radius: 4px;
}

/* ================== LABEL ================== */
QLabel {
    background: transparent;
    border: none;
}

QLabel[role="week-label"] {
    font-size: 20pt;
    font-weight: bold;
}

/* ================== КНОПКИ ================== */
QPushButton {
    background-color: #4a86c5;
    color: #ffffff;
    border: none;
    padding: 6px 12px;
    border-radius: 4px;
}

QPushButton:hover {
    background-color: #5a96d5;
}

QPushButton:pressed {
    background-color: #3a6ea5;
}

QPushButton[role="week-nav"] {
    background-color: #4a86c5;
    border: none;
    color: #a6c8ff;
    font-size: 20pt;
    font-weight: bold;
}

QPushButton[role="week-nav"]:hover {
    background-color: #4a86c5;
}

QPushButton[role="week-nav"]:pressed {
    background-color: rgba(255, 255, 255, 0.15);
}
QToolButton#viewSwitchBtn {
    background: transparent;
    border: none;
    padding: 4px;
    font-weight: bold;
}

QToolButton#viewSwitchBtn:hover {
    background-color: rgba(90,150,213,0.15);
    border-radius: 12px;
}

QToolButton#viewSwitchBtn:checked {
    background-color: rgba(90,150,213,0.30);
    border-radius: 14px;
}

/* ================== TOOL BUTTON ================== */
QToolButton {
    background-color: transparent;
    border: none;
    padding: 4px;
}

QToolButton:hover {
    background-color: #3d5a73;
}

/* ================== COMBOBOX ================== */
QComboBox {
    background-color: #2f3133;
    border: 1px solid #555555;
    padding: 4px;
    border-radius: 4px;
    min-height: 18px;
}

QComboBox QAbstractItemView {
    background-color: #2f3133;
    selection-background-color: #3d5a73;
}

/* ================== RADIO / CHECK ================== */
QRadioButton, QCheckBox {
    spacing: 6px;
    font-weight: bold;
}
QRadioButton:disabled {
    color: #777777;
}

QRadioButton::indicator:disabled {
    background-color: #555555;
    border: 1px solid #444444;
    width: 14px;
    height: 14px;
    border-radius: 7px;
}

/* ================== ТАБЛИЦА ================== */
QTableView {
    background-color: #2f3133;
    gridline-color: #444444;
    selection-background-color: #3d5a73;
    selection-color: #ffffff;
    alternate-background-color: #2a2c2e;
}

QTableView::item {
    padding: 4px;
}

QTableView::item:selected {
    background-color: #3d5a73;
}

/* ================== ЗАГОЛОВКИ ТАБЛИЦЫ ================== */
QHeaderView::section {
    background-color: #3a3d41;
    border: 1px solid #444444;
    padding: 6px;
    font-weight: bold;
}

/* ================== SCROLLBAR ================== */
QScrollBar:vertical {
    background: #2b2b2b;
    width: 12px;
}

QScrollBar::handle:vertical {
    background: #555555;
    min-height: 20px;
    border-radius: 6px;
}

QScrollBar::handle:vertical:hover {
    background: #666666;
}

/* ================== TEXT EDIT (детализация) ================== */
QTextEdit {
    background-color: #2f3133;
    border: 1px solid #444444;
    border-radius: 4px;
    padding: 6px;
}

/* ================== SPLITTER ================== */
QSplitter::handle {
    background-color: #444444;
}
/* ================== QGroupBox ================== */
QGroupBox {
    font-weight: bold;
    border: 1px solid #444444;
    border-radius: 6px;
    margin-top: 8px;
}

QGroupBox::title {
    subcontrol-origin: margin;
    left: 10px;
    padding: 0 4px;
}
/* ================== DateEdit ================== */
QDateEdit {
    background-color: #2f3133;
    border: 1px solid #555555;
    border-radius: 4px;
    padding: 4px 6px;
    min-height: 18px;
    color: #e6e6e6;
}

QDateEdit:hover {
    border: 1px solid #5a96d5;
}

QDateEdit:focus {
    border: 1px solid #4a86c5;
}

QDateEdit::down-arrow {
    width: 10px;
    height: 10px;
}
"""


def excepthook(type, value, tb):
    print("UNCAUGHT EXCEPTION:")
    traceback.print_exception(type, value, tb)

def main():
    app = QApplication(sys.argv)

    script_dir = os.path.dirname(os.path.abspath(__file__))
    icon_path = os.path.join(script_dir, "Case_analysis.ico")
    app.setWindowIcon(QIcon(icon_path))

    app.setStyleSheet(LIGHT_STYLE)

    window = MainWindow()
    window.resize(1200, 800)
    window.showMaximized()
    window.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    sys.excepthook = excepthook
    main()
